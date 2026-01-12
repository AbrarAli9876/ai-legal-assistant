import httpx
from dotenv import load_dotenv
from docx import Document
import io

from app.core.config import GENAI_API_KEY

# --- Configuration ---

GEMINI_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-preview-09-2025:generateContent"
API_KEY = GENAI_API_KEY

# System instruction for the AI document drafter
SYSTEM_INSTRUCTION = """You are an expert AI legal assistant specializing in drafting clear, professional, and accurate legal documents based on Indian law.
-   **Task:** Generate the full text for the requested legal document.
-   **Input:** You will receive the document type and key details (like party names, dates, terms).
-   **Output:** Return ONLY the complete, professionally formatted text of the document.
-   **Formatting:** Use standard legal document formatting. Use newlines (\n) for paragraphs and numbering.
-   **Tone:** Formal, precise, and authoritative.
-   **CRITICAL:** Do NOT include any conversational text, disclaimers, or explanations. Only output the raw document text itself.
"""

# --- Helper Function: Call Gemini API ---

async def _call_gemini_for_drafting(prompt: str) -> str:
    """Helper to call Gemini API for text generation."""
    if not API_KEY:
        raise ValueError("Server configuration error: GENAI_API_KEY not found.")

    full_api_url = f"{GEMINI_API_URL}?key={API_KEY}"
    
    payload = {
        "contents": [{"parts": [{"text": prompt}]}],
        "systemInstruction": {"parts": [{"text": SYSTEM_INSTRUCTION}]},
        "generationConfig": {
            "temperature": 0.5,
            "topK": 1,
            "topP": 1,
            "maxOutputTokens": 4096,
        }
    }

    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(full_api_url, json=payload, headers={"Content-Type": "application/json"})
            response.raise_for_status()
            
            result = response.json()
            text = result.get("candidates", [{}])[0].get("content", {}).get("parts", [{}])[0].get("text")
            
            if not text:
                raise ValueError("AI failed to generate document text.")
            return text
            
    except httpx.HTTPStatusError as e:
        print(f"HTTP error occurred: {e}")
        raise ValueError(f"AI API request failed: {e.response.text}")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        raise

# --- Main Service Function ---

async def create_document(doc_type: str, details: dict) -> io.BytesIO:
    """
    Generates a legal document text using AI and creates a .docx file in memory.
    """
    
    # 1. Create a detailed prompt for the AI
    # This can be expanded with more logic for each doc_type
    prompt = f"Generate a {doc_type} with the following details:\n"
    for key, value in details.items():
        if value: # Only include details that are filled in
            prompt += f"- {key.replace('_', ' ').title()}: {value}\n"
    
    prompt += "\nReturn only the full, formatted text of the legal document, ready for a .docx file."

    try:
        # 2. Get the raw document text from the AI
        print(f"Generating draft for: {doc_type}")
        document_text = await _call_gemini_for_drafting(prompt)
        print("Draft received from AI.")
        
        # 3. Create a .docx file in memory
        document = Document()
        
        # Add a title (optional, but good practice)
        document.add_heading(doc_type, level=0)
        
        # Add the AI-generated text
        # We split by newline to preserve paragraph breaks
        for para in document_text.split('\n'):
            if para.strip(): # Avoid adding empty paragraphs
                document.add_paragraph(para)
        
        # 4. Save the document to a byte stream
        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0) # Rewind the stream to the beginning
        
        print(f"Successfully created .docx for {doc_type} in memory.")
        return file_stream
        
    except Exception as e:
        print(f"Error in create_document service: {e}")
        raise